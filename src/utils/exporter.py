import re
from io import BytesIO
from docx import Document
from docx.shared import Pt

def convert_to_docx(markdown_text: str) -> BytesIO:
  """
  Convert Markdown content to a DOCX file in memory.
  
  Args:
    markdown_text: The SRS content in Markdown format
      
  Returns:
    BytesIO: A binary stream containing the DOCX file
  """
  document = Document()
  
  # Set default style
  style = document.styles['Normal']
  font = style.font
  font.name = 'Calibri'
  font.size = Pt(11)
  
  lines = markdown_text.split('\n')
  
  in_code_block = False
  
  for line in lines:
    line = line.rstrip()
    
    # Handle code blocks
    if line.startswith('```'):
      in_code_block = not in_code_block
      continue
        
    if in_code_block:
      p = document.add_paragraph()
      runner = p.add_run(line)
      runner.font.name = 'Courier New'
      runner.font.size = Pt(10)
      continue
        
    # Headers
    if line.startswith('# '):
      document.add_heading(line[2:], level=1)
    elif line.startswith('## '):
      document.add_heading(line[3:], level=2)
    elif line.startswith('### '):
      document.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
      document.add_heading(line[5:], level=4)
        
    # Bullet points
    elif line.strip().startswith('- ') or line.strip().startswith('* '):
      p = document.add_paragraph(line.strip()[2:], style='List Bullet')
        
    # Numbered lists
    elif re.match(r'^\d+\. ', line.strip()):
      content = re.sub(r'^\d+\. ', '', line.strip())
      p = document.add_paragraph(content, style='List Number')
        
    # Horizontal Rule
    elif line.strip() == '---' or line.strip() == '***':
      document.add_paragraph('_' * 20)
        
    # Ignore empty lines unless inside code block
    elif not line.strip():
      continue
        
    # Normal text (handling bold/italic crudely)
    else:
      p = document.add_paragraph(line)

  # Save to memory
  file_stream = BytesIO()
  document.save(file_stream)
  file_stream.seek(0)
  
  return file_stream
